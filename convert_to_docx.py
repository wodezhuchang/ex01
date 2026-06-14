# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import re

def set_default_font(document):
    for style in document.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style.font.name = 'SimSun'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            style.font.size = Pt(10.5)

def add_runs_with_formatting(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            inner_text = part[2:-2]
            if inner_text:
                run = paragraph.add_run(inner_text)
                run.font.name = 'SimSun'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                run.font.size = Pt(10.5)
                run.font.bold = True
        else:
            run = paragraph.add_run(part)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(10.5)

def add_cover(document):
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run('\n\n\n成绩管理系统实现报告')
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    for _ in range(8):
        document.add_paragraph()

    info = [
        ('课程名称：', '________________________'),
        ('姓名：', '________________________'),
        ('学号：', '________________________'),
        ('班级：', '________________________'),
        ('提交日期：', '______年______月______日'),
        ('指导教师：', '________________________'),
    ]

    for label, value in info:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(label)
        run_label.font.name = 'SimHei'
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        run_label.font.size = Pt(18)
        run_label.font.bold = True
        run_label.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
        
        run_value = p.add_run(value)
        run_value.font.name = 'Times New Roman'
        run_value.font.size = Pt(18)
        run_value.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    document.add_page_break()

def add_heading(document, text, level=1):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

def add_text(document, text, bold=False, italic=False, underline=False):
    p = document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_runs_with_formatting(p, text)
    return p

def add_bullet(document, text, level=0):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    bullet_char = '● ' if level == 0 else '○ '
    run = p.add_run(bullet_char)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(10.5)
    add_runs_with_formatting(p, text)
    return p

def add_numbered_list(document, text, number):
    p = document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(f'{number}. ')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(10.5)
    add_runs_with_formatting(p, text)
    return p

def add_code_block(document, code_text):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p

def add_table_from_markdown(document, markdown_table):
    lines = markdown_table.strip().split('\n')
    if len(lines) < 3:
        return None

    header_line = lines[0]
    data_lines = lines[2:]

    headers = [h.strip().replace('**', '') for h in header_line.strip('|').split('|')]
    rows = []
    for line in data_lines:
        if line.strip():
            cells = [c.strip().replace('**', '') for c in line.strip('|').split('|')]
            rows.append(cells)

    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.bold = True
                run.font.size = Pt(10.5)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'SimSun'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                    run.font.size = Pt(10.5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    return table

def add_horizontal_rule(document):
    p = document.add_paragraph()
    run = p.add_run('━' * 80)
    run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)
    run.font.size = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def parse_markdown_to_docx(md_content, docx_path):
    document = Document()

    sections = document.sections
    for section in sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    set_default_font(document)

    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    in_table = False
    table_content = []

    while i < len(lines):
        line = lines[i]
        stripped_line = line.strip()

        if stripped_line.startswith('```'):
            if in_code_block:
                full_code = '\n'.join(code_content)
                add_code_block(document, full_code)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
                code_content = []
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        if stripped_line.startswith('|') and '|' in stripped_line:
            if not in_table:
                in_table = True
                table_content = [stripped_line]
            else:
                table_content.append(stripped_line)
            i += 1
            continue
        elif in_table:
            full_table = '\n'.join(table_content)
            add_table_from_markdown(document, full_table)
            in_table = False
            table_content = []

        if stripped_line == '---':
            add_horizontal_rule(document)
            i += 1
            continue

        if stripped_line.startswith('# '):
            add_heading(document, stripped_line[2:].replace('**', ''), level=1)
        elif stripped_line.startswith('## '):
            add_heading(document, stripped_line[3:].replace('**', ''), level=1)
        elif stripped_line.startswith('### '):
            add_heading(document, stripped_line[4:].replace('**', ''), level=2)
        elif stripped_line.startswith('#### '):
            add_heading(document, stripped_line[5:].replace('**', ''), level=3)
        elif re.match(r'^\d+\.\s+', stripped_line):
            match = re.match(r'^(\d+)\.\s+(.+)$', stripped_line)
            if match:
                num = match.group(1)
                text = match.group(2)
                add_numbered_list(document, text, num)
        elif stripped_line.startswith('- '):
            add_bullet(document, stripped_line[2:])
        elif stripped_line.startswith('   - '):
            add_bullet(document, stripped_line[5:], level=1)
        elif stripped_line.startswith('      - '):
            add_bullet(document, stripped_line[8:], level=2)
        elif stripped_line:
            p = document.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            add_runs_with_formatting(p, stripped_line)
        elif stripped_line == '' and i > 0 and i < len(lines) - 1:
            pass

        i += 1

    if in_table:
        full_table = '\n'.join(table_content)
        add_table_from_markdown(document, full_table)

    document.save(docx_path)
    print(f'Document saved to: {docx_path}')

if __name__ == '__main__':
    import sys
    import os

    if len(sys.argv) >= 3:
        md_file = sys.argv[1]
        docx_file = sys.argv[2]
    else:
        md_file = r'c:\Users\yangd\Documents\GitHub\ex01\图书管理系统实验报告.md'
        docx_file = r'c:\Users\yangd\Documents\GitHub\ex01\图书管理系统实验报告.docx'

    if not os.path.exists(md_file):
        print(f'错误：找不到 Markdown 文件: {md_file}')
        sys.exit(1)

    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    parse_markdown_to_docx(md_content, docx_file)
